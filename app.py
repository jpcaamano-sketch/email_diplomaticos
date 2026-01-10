import streamlit as st
import google.generativeai as genai
from fpdf import FPDF
from docx import Document
import io
import re

# --- 1. CONFIGURACIÓN ---
st.set_page_config(page_title="Traductor Diplomático", layout="centered")

st.markdown("""
<style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .block-container {padding-top: 2rem;}
    .stAlert { margin-bottom: 1rem; }
</style>
""", unsafe_allow_html=True)

try:
    api_key = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=api_key)
except Exception:
    st.error("⚠️ Falta configurar la API Key en .streamlit/secrets.toml")
    st.stop()

# --- 2. LÓGICA IA (ROBUSTA CON REGEX) ---
def generar_opciones(texto, destinatario):
    try:
        # Usamos el modelo rápido y estable
        model = genai.GenerativeModel("gemini-1.5-flash")
        
        prompt = f"""
        TAREA: Actúa como experto en comunicación corporativa.
        OBJETIVO: Reescribir el siguiente texto borrador para un "{destinatario}".
        
        TEXTO ORIGINAL: "{texto}"
        
        INSTRUCCIONES:
        1. No preguntes nada. Si el texto es corto, interprétalo y mejóralo.
        2. Genera 3 versiones obligatoriamente.
        3. Usa EXACTAMENTE estos títulos para separar las versiones:
        
        SECCION_PROFESIONAL:
        [Versión formal y educada aquí]
        
        SECCION_DIRECTA:
        [Versión ejecutiva y al grano aquí]
        
        SECCION_COLOQUIAL:
        [Versión cercana y amable aquí]
        """
        
        response = model.generate_content(prompt)
        resultado = response.text
        
        # Búsqueda inteligente (Regex) para evitar errores si el modelo habla de más
        prof = re.search(r"SECCION_PROFESIONAL:(.*?)(?=SECCION_DIRECTA:|$)", resultado, re.DOTALL | re.IGNORECASE)
        directa = re.search(r"SECCION_DIRECTA:(.*?)(?=SECCION_COLOQUIAL:|$)", resultado, re.DOTALL | re.IGNORECASE)
        coloquial = re.search(r"SECCION_COLOQUIAL:(.*?)(?=$)", resultado, re.DOTALL | re.IGNORECASE)
        
        return {
            "profesional": prof.group(1).strip() if prof else "No se pudo generar.",
            "directo": directa.group(1).strip() if directa else "No se pudo generar.",
            "coloquial": coloquial.group(1).strip() if coloquial else "No se pudo generar."
        }

    except Exception as e:
        return {"error": str(e)}

# --- 3. GENERADORES DE ARCHIVOS ---
def generar_archivo(resultados, original, formato):
    if formato == "Word (.docx)":
        doc = Document()
        doc.add_heading('Propuestas de Comunicación', 0)
        doc.add_heading('Original:', level=2)
        doc.add_paragraph(original)
        
        doc.add_heading('1. Profesional', level=1)
        doc.add_paragraph(resultados.get('profesional', ''))
        
        doc.add_heading('2. Directo', level=1)
        doc.add_paragraph(resultados.get('directo', ''))
        
        doc.add_heading('3. Coloquial', level=1)
        doc.add_paragraph(resultados.get('coloquial', ''))
        
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
        
    else: # PDF
        pdf = FPDF()
        pdf.add_page()
        
        def L(t): return t.encode('latin-1', 'replace').decode('latin-1') if t else ""
        
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, L('Propuestas de Comunicación'), 0, 1, 'C')
        pdf.ln(5)
        
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "Original:", 0, 1)
        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 6, L(original))
        pdf.ln(5)
        
        orden = [("1. Profesional", "profesional"), ("2. Directo", "directo"), ("3. Coloquial", "coloquial")]
        for tit, key in orden:
            pdf.set_font("Arial", 'B', 12)
            pdf.set_text_color(0, 50, 100)
            pdf.cell(0, 10, L(tit), 0, 1)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", size=11)
            pdf.multi_cell(0, 6, L(resultados.get(key, '')))
            pdf.ln(3)
            
        return pdf.output(dest='S').encode('latin-1'), "application/pdf", "pdf"

# --- 4. INTERFAZ VISUAL ---
st.title("🕊️ Traductor Diplomático")
st.caption("Convierte borradores difíciles en comunicación efectiva.")
st.divider()

# 1. INPUTS
destinatario = st.selectbox("1. ¿A quién le escribes?", 
    ["Cliente", "Jefe/Superior", "Colaborador/Equipo", "Proveedor", "Par (Colega/Igual)"])

texto_input = st.text_area("2. Borrador del texto (sin filtro):", height=120, 
    placeholder="Ej: Necesito que me entregues eso ahora mismo...")

# Estado de sesión
if 'resultado_v3' not in st.session_state:
    st.session_state.resultado_v3 = None

# Botón de Acción
if st.button("✨ Generar Propuestas", type="primary", use_container_width=True):
    if not texto_input:
        st.warning("Escribe un borrador primero.")
    else:
        with st.spinner("Redactando versiones..."):
            st.session_state.resultado_v3 = generar_opciones(texto_input, destinatario)

# 2. RESULTADOS
if st.session_state.resultado_v3:
    res = st.session_state.resultado_v3
    
    if "error" in res:
        st.error(f"Error técnico: {res['error']}")
    else:
        st.markdown("### 📢 Opciones Asertivas")
        
        # Tarjetas visuales
        st.info(f"**👔 Profesional (Formal):**\n\n{res.get('profesional')}")
        st.warning(f"**⚡ Directo (Ejecutivo):**\n\n{res.get('directo')}")
        st.success(f"**☕ Coloquial (Cercano):**\n\n{res.get('coloquial')}")
        
        st.divider()
        
        # 3. ZONA DE DESCARGA
        st.subheader("📥 Descargar Archivo")
        
        col_name, col_type = st.columns([2, 1])
        with col_name:
            nombre_archivo = st.text_input("Nombre del archivo:", value="Mis_Propuestas", help="Sin extensión")
        with col_type:
            tipo_archivo = st.radio("Formato:", ["Word (.docx)", "PDF (.pdf)"], horizontal=True)
            
        # Preparar el archivo
        bytes_data, mime, ext = generar_archivo(res, texto_input, tipo_archivo)
        
        st.download_button(
            label=f"💾 Bajar en {tipo_archivo}",
            data=bytes_data,
            file_name=f"{nombre_archivo}.{ext}",
            mime=mime,
            use_container_width=True
        )
